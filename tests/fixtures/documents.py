"""Sample clinical documents.

Built programmatically rather than committed as binary fixtures, for three reasons: a
committed PDF containing realistic clinical text is a PHI-shaped artifact nobody wants in
version control, generated documents can be parameterised per test, and a binary blob
gives a reviewer no way to see what is being tested.

All patient details are fabricated.
"""

from __future__ import annotations

import io

import pytest

__all__ = [
    "DISCHARGE_SUMMARY_TEXT",
    "LAB_REPORT_TEXT",
    "RADIOLOGY_NOTE_TEXT",
    "build_docx",
    "build_pdf",
    "build_scanned_pdf",
]

DISCHARGE_SUMMARY_TEXT = """DISCHARGE SUMMARY

Patient Name: Jordan Rivera
MRN: 00471925
Date of Service: 2026-03-14

CHIEF COMPLAINT
Substernal chest pain radiating to the left arm.

HISTORY OF PRESENT ILLNESS
The patient is a 62-year-old presenting with a three day history of exertional chest
discomfort. Symptoms worsened overnight prior to admission. Pt. was given 2.5 mg of
morphine q.i.d. per Dr. Chen. No prior cardiac history is documented.

PAST MEDICAL HISTORY
Hypertension. Type 2 diabetes mellitus. Hyperlipidemia.

ALLERGIES
Penicillin - rash.

MEDICATIONS
Lisinopril | 10 mg | daily
Metformin | 500 mg | twice daily
Atorvastatin | 40 mg | nightly

LABORATORY RESULTS
Troponin I | 0.04 ng/mL | reference range 0.00-0.04
Hemoglobin | 13.8 g/dL | reference range 13.5-17.5

ASSESSMENT AND PLAN
Acute coronary syndrome ruled out. Continue medical management and outpatient follow up.

HOSPITAL COURSE
The patient remained hemodynamically stable throughout the admission and was ambulating
without symptoms by hospital day two.

DISPOSITION
Discharged home in stable condition.
"""

LAB_REPORT_TEXT = """LABORATORY REPORT

Specimen: Serum
Collected: 2026-02-01
Reported: 2026-02-02

RESULTS
Sodium | 141 | mmol/L | 135-145
Potassium | 4.2 | mmol/L | 3.5-5.1
Creatinine | 1.02 | mg/dL | 0.70-1.30
White blood cell | 7.4 | K/uL | 4.0-11.0
Hemoglobin | 14.1 | g/dL | 13.5-17.5
"""

RADIOLOGY_NOTE_TEXT = """RADIOLOGY REPORT

TECHNIQUE
Non-contrast CT scan of the chest was performed.

COMPARISON
Prior radiograph dated 2025-11-04.

FINDINGS
No focal consolidation. No pleural effusion. Cardiomediastinal silhouette is within
normal limits.

IMPRESSION
No acute cardiopulmonary process.
"""


def build_pdf(text: str, *, title: str | None = None, pages: int | None = None) -> bytes:
    """Render text into a real, parseable PDF with a genuine text layer."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    if title:
        pdf.setTitle(title)

    _width, height = LETTER
    margin = 60
    line_height = 14
    max_lines = int((height - 2 * margin) / line_height)

    lines = text.split("\n")
    chunks = [lines[i : i + max_lines] for i in range(0, len(lines), max_lines)] or [[""]]
    if pages is not None:
        chunks = chunks[:pages] if len(chunks) >= pages else chunks + [[""]] * (pages - len(chunks))

    for page_lines in chunks:
        pdf.setFont("Helvetica", 10)
        y = height - margin
        for line in page_lines:
            pdf.drawString(margin, y, line[:110])
            y -= line_height
        pdf.showPage()

    pdf.save()
    return buffer.getvalue()


def build_scanned_pdf(*, pages: int = 1) -> bytes:
    """Build a PDF with no text layer, standing in for a scanned document.

    Each page draws only a filled rectangle, so text extraction legitimately returns
    nothing and the OCR routing path is exercised for real rather than simulated.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    for _ in range(pages):
        pdf.setFillColorRGB(0.85, 0.85, 0.85)
        pdf.rect(72, 72, 400, 600, fill=1, stroke=0)
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def build_docx(text: str, *, title: str | None = None, author: str | None = None) -> bytes:
    """Build a DOCX with heading styles and a real table.

    Lines that are entirely uppercase become ``Heading 1`` paragraphs and pipe-delimited
    lines become table rows, so the parser's style- and table-handling paths are covered.
    """
    import docx

    document = docx.Document()
    if title:
        document.core_properties.title = title
    if author:
        document.core_properties.author = author

    pending_rows: list[list[str]] = []

    def _flush_table() -> None:
        if not pending_rows:
            return
        table = document.add_table(rows=len(pending_rows), cols=len(pending_rows[0]))
        for row_index, row in enumerate(pending_rows):
            for cell_index, value in enumerate(row):
                table.rows[row_index].cells[cell_index].text = value
        pending_rows.clear()

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            _flush_table()
            continue
        if "|" in line:
            pending_rows.append([cell.strip() for cell in line.split("|")])
            continue
        _flush_table()
        if line.isupper() and len(line) <= 80:
            document.add_heading(line, level=1)
        else:
            document.add_paragraph(line)
    _flush_table()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def discharge_summary_text() -> bytes:
    return DISCHARGE_SUMMARY_TEXT.encode("utf-8")


@pytest.fixture
def discharge_summary_pdf() -> bytes:
    return build_pdf(DISCHARGE_SUMMARY_TEXT, title="Discharge Summary")


@pytest.fixture
def discharge_summary_docx() -> bytes:
    return build_docx(DISCHARGE_SUMMARY_TEXT, title="Discharge Summary", author="Test Author")


@pytest.fixture
def lab_report_text() -> bytes:
    return LAB_REPORT_TEXT.encode("utf-8")


@pytest.fixture
def radiology_note_text() -> bytes:
    return RADIOLOGY_NOTE_TEXT.encode("utf-8")


@pytest.fixture
def scanned_pdf() -> bytes:
    return build_scanned_pdf()
