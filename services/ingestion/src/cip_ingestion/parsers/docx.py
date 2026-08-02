"""DOCX parser.

DOCX carries explicit structure that PDF only implies: paragraph styles name headings,
tables are real tables, and core properties hold authored metadata. The parser uses those
signals directly rather than re-deriving them heuristically — a ``Heading 2`` style is
ground truth in a way that "line is short and uppercase" never is.

Two structural details of the format shape the implementation. First, ``python-docx``
exposes paragraphs and tables through separate collections, losing their relative order;
the underlying XML body preserves it, so the parser walks the body elements to keep a
table adjacent to the paragraph that introduces it. Second, DOCX has no page concept —
pagination is a rendering-time property of the consuming application — so the whole
document is emitted as a single logical page rather than fabricating page numbers that
would produce false citations.
"""

from __future__ import annotations

from typing import Any

from cip_core.logging import get_logger
from cip_ingestion.domain import BlockKind, ParsedDocument, ParsedPage, TextBlock
from cip_ingestion.parsers.base import ParserError

__all__ = ["DocxParser"]

_log = get_logger(__name__)

_DOCX_MEDIA_TYPES = frozenset(
    {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
)


class DocxParser:
    """Extracts text, headings, and tables from DOCX bytes."""

    @property
    def name(self) -> str:
        return "docx"

    @property
    def supported_media_types(self) -> frozenset[str]:
        return _DOCX_MEDIA_TYPES

    def parse(self, data: bytes, *, filename: str | None = None) -> ParsedDocument:
        import io

        import docx
        from docx.opc.exceptions import PackageNotFoundError

        if not data:
            raise ParserError("DOCX payload is empty", media_type=next(iter(_DOCX_MEDIA_TYPES)))

        try:
            document = docx.Document(io.BytesIO(data))
        except PackageNotFoundError as exc:
            raise ParserError(
                "File is not a valid DOCX package", media_type=next(iter(_DOCX_MEDIA_TYPES))
            ) from exc
        except Exception as exc:
            raise ParserError(
                f"DOCX could not be opened: {type(exc).__name__}",
                media_type=next(iter(_DOCX_MEDIA_TYPES)),
            ) from exc

        warnings: list[str] = []
        blocks: list[TextBlock] = []
        try:
            blocks = list(self._iter_body_blocks(document))
        except Exception as exc:
            warnings.append(f"body traversal failed ({type(exc).__name__}); used flat extraction")
            blocks = list(self._iter_paragraphs_flat(document))

        properties = self._extract_properties(document)

        return ParsedDocument(
            parser_name=self.name,
            media_type=next(iter(_DOCX_MEDIA_TYPES)),
            pages=(ParsedPage(page_number=1, blocks=tuple(blocks)),),
            properties=properties,
            warnings=tuple(warnings),
        )

    def _iter_body_blocks(self, document: Any) -> Any:
        """Walk body XML so paragraphs and tables keep their document order."""
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = document.element.body
        order = 0
        for child in body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if not text:
                    continue
                yield TextBlock(
                    text=text, kind=self._paragraph_kind(paragraph), page_number=1, order=order
                )
                order += 1
            elif tag == "tbl":
                rendered = self._render_table(Table(child, document))
                if rendered:
                    yield TextBlock(text=rendered, kind=BlockKind.TABLE, page_number=1, order=order)
                    order += 1

    def _iter_paragraphs_flat(self, document: Any) -> Any:
        """Fallback that ignores ordering but still recovers paragraph text."""
        order = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            yield TextBlock(
                text=text, kind=self._paragraph_kind(paragraph), page_number=1, order=order
            )
            order += 1

    @staticmethod
    def _paragraph_kind(paragraph: Any) -> BlockKind:
        """Map a paragraph style onto a block kind.

        Style names are matched case-insensitively by prefix because Word localises them
        and templates add suffixes (``Heading 1``, ``heading1``, ``Heading 1 Char``).
        """
        style_name = (getattr(paragraph.style, "name", "") or "").strip().lower()
        if style_name.startswith(("heading", "title", "subtitle")):
            return BlockKind.HEADING
        if style_name.startswith(("list", "bullet")):
            return BlockKind.LIST_ITEM
        return BlockKind.PARAGRAPH

    @staticmethod
    def _render_table(table: Any) -> str:
        """Render a table as pipe-delimited rows, matching the PDF parser's convention."""
        lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        return "\n".join(lines)

    @staticmethod
    def _extract_properties(document: Any) -> dict[str, Any]:
        properties: dict[str, Any] = {}
        core = getattr(document, "core_properties", None)
        if core is None:
            return properties
        for attribute, key in (
            ("title", "title"),
            ("author", "author"),
            ("subject", "subject"),
            ("last_modified_by", "last_modified_by"),
        ):
            value = getattr(core, attribute, None)
            if isinstance(value, str) and value.strip():
                properties[key] = value.strip()[:512]
        created = getattr(core, "created", None)
        if created is not None:
            properties["creation_date"] = created.isoformat()
        return properties
